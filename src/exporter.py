# app/exporter.py

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Any

def _ensure_styles(doc: Document):
    """
    Crea estilos personalizados si no existen:
    - FunctionHeading: Heading nivel 2 con tamaño 14pt
    - ClassHeading: Heading nivel 2 con tamaño 14pt y color
    - CodeStyle: Fuente monoespacio 10pt
    """
    styles = doc.styles

    if 'FunctionHeading' not in styles:
        func_h = styles.add_style('FunctionHeading', WD_STYLE_TYPE.PARAGRAPH)
        func_h.base_style = styles['Heading 2']
        func_h.font.size = Pt(14)

    if 'ClassHeading' not in styles:
        class_h = styles.add_style('ClassHeading', WD_STYLE_TYPE.PARAGRAPH)
        class_h.base_style = styles['Heading 2']
        class_h.font.size = Pt(14)
        class_h.font.color.rgb = RGBColor(0x00, 0x30, 0x60)

    if 'CodeStyle' not in styles:
        code_s = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
        code_s.font.name = 'Courier New'
        code_s.font.size = Pt(10)

def _add_cover_page(doc: Document, title: str):
    """Agrega una portada con título centrado y salto de página."""
    p = doc.add_paragraph()
    p.alignment = 1  # centrado
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    doc.add_page_break()

def _add_function_block(doc: Document, block: Dict[str, Any]):
    """Genera sección para una función."""
    doc.add_paragraph(f"Función: {block['name']}", style='FunctionHeading')
    doc.add_paragraph(f"Docstring: {block.get('docstring', '—')}")
    doc.add_paragraph(f"Argumentos: {', '.join(block.get('args', [])) or '—'}")

    doc.add_paragraph("Lógica:", style='List Bullet')
    for logic in block.get('logic', []):
        if logic['type'] == 'return':
            p = doc.add_paragraph(f"Retorna: {logic['value']}", style='Intense Quote')
            p.style = 'CodeStyle'
        elif logic['type'] == 'comment':
            doc.add_paragraph(f"Comentario: {logic['value']}", style='Quote')

def _add_class_block(doc: Document, block: Dict[str, Any]):
    """Genera sección para una clase y sus métodos."""
    doc.add_paragraph(f"Clase: {block['name']}", style='ClassHeading')
    doc.add_paragraph(f"Docstring: {block.get('docstring', '—')}")

    for method in block.get('methods', []):
        doc.add_paragraph(f"Método: {method['name']}", style='List Number')
        doc.add_paragraph(f"Docstring: {method.get('docstring', '—')}")
        doc.add_paragraph(f"Argumentos: {', '.join(method.get('args', [])) or '—'}")

def generate_doc(parsed_blocks: List[Dict[str, Any]],
    title: str = "Documentación Generada") -> Document:
    """
    Función principal que construye el documento.
    - parsed_blocks: lista de bloques parseados por tipo.
    - title: título que irá en la portada.
    """
    doc = Document()
    _ensure_styles(doc)
    _add_cover_page(doc, title)

    for block in parsed_blocks:
        btype = block.get('type')
        if btype == 'function':
            _add_function_block(doc, block)
        elif btype == 'class':
            _add_class_block(doc, block)
        doc.add_page_break()

    return doc
